"""
Brand Registration API Module
Handles brand registration data including file uploads and settings
"""

import os
import json
import time
import shutil
import re
from pathlib import Path
from typing import Dict, Any, Optional

from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from pydantic import BaseModel
from datetime import datetime

try:
    # Lightweight PDF parsing for text extraction
    from PyPDF2 import PdfReader  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    PdfReader = None  # type: ignore

try:
    # Optional DOCX support
    import docx  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    docx = None  # type: ignore

try:
    # OCR support for scanned PDFs
    import pytesseract  # type: ignore
    from pdf2image import convert_from_path  # type: ignore
    OCR_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency
    OCR_AVAILABLE = False
    pytesseract = None  # type: ignore

# Create router for brand registration endpoints
router = APIRouter(prefix="/api/brand-registration", tags=["brand-registration"])

# Project root for absolute paths (production-safe)
_ROOT = Path(os.environ.get("PROJECT_ROOT", str(Path(__file__).resolve().parent.parent)))

# Data models
class BrandSettings(BaseModel):
    defaultLanguage: str = "en"
    defaultLLM: str = "gpt-4.1"
    productDefaultPct: int = 30

class VoiceProfile(BaseModel):
    formality: int = 30
    humor: int = 60
    warmth: int = 70
    emojiPolicy: str = "medium"

class ContentPillar(BaseModel):
    name: str
    description: str
    weight: int

class BrandPolicies(BaseModel):
    forbiddenPhrases: list[str] = []
    maxHashtags: int = 5
    brandHashtags: list[str] = []

class BrandBlueprint(BaseModel):
    version: str = "1.0.0"
    status: str = "draft"
    voice: VoiceProfile = VoiceProfile()
    pillars: list[ContentPillar] = []
    policies: BrandPolicies = BrandPolicies()
    productDefaultPct: int = 30

class BrandRegistrationData(BaseModel):
    id: str
    name: str
    guidelineDoc: Dict[str, Any] = {"name": "", "size": 0, "status": "none"}
    blueprint: BrandBlueprint = BrandBlueprint()
    settings: BrandSettings = BrandSettings()
    createdAt: str
    updatedAt: str

class BrandRegistrationService:
    """Service class to handle brand registration operations"""
    
    def __init__(self, data_file: Optional[str] = None, upload_dir: Optional[Path] = None, root: Optional[Path] = None):
        root = root or _ROOT
        self.data_file = Path(data_file) if data_file else (root / "brand_registrations.json")
        self.upload_dir = upload_dir if upload_dir is not None else (root / "brand registration" / "uploads")
        self.upload_dir = Path(self.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
    def _load_data(self) -> Dict[str, Any]:
        """Load brand registration data from JSON file"""
        if self.data_file.exists():
            try:
                with open(self.data_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except (json.JSONDecodeError, IOError):
                return {"brands": {}}
        return {"brands": {}}
    
    def _save_data(self, data: Dict[str, Any]) -> None:
        """Save brand registration data to JSON file"""
        try:
            with open(self.data_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except IOError as e:
            raise HTTPException(status_code=500, detail=f"Failed to save data: {str(e)}")
    
    def create_brand(self, brand_data: BrandRegistrationData) -> Dict[str, Any]:
        """Create a new brand registration"""
        data = self._load_data()
        
        # Generate unique ID if not provided
        if not brand_data.id:
            brand_data.id = f"brand-{int(time.time())}"
        
        # Set timestamps
        now = datetime.now().isoformat()
        brand_data.createdAt = now
        brand_data.updatedAt = now
        
        # Save to data store
        data["brands"][brand_data.id] = brand_data.dict()
        self._save_data(data)
        
        return {"success": True, "brand_id": brand_data.id, "data": brand_data.dict()}
    
    def update_brand(self, brand_id: str, updates: Dict[str, Any]) -> Dict[str, Any]:
        """Update existing brand registration"""
        data = self._load_data()
        
        if brand_id not in data["brands"]:
            raise HTTPException(status_code=404, detail="Brand not found")
        
        # Update the brand data
        brand_data = data["brands"][brand_id]
        brand_data.update(updates)
        brand_data["updatedAt"] = datetime.now().isoformat()
        
        self._save_data(data)
        
        return {"success": True, "brand_id": brand_id, "data": brand_data}
    
    def get_brand(self, brand_id: str) -> Dict[str, Any]:
        """Get brand registration by ID"""
        data = self._load_data()
        
        if brand_id not in data["brands"]:
            raise HTTPException(status_code=404, detail="Brand not found")
        
        return data["brands"][brand_id]
    
    def list_brands(self) -> Dict[str, Any]:
        """List all brand registrations"""
        data = self._load_data()
        return {"brands": list(data["brands"].values())}
    
    def save_uploaded_file(self, file: UploadFile, brand_id: str) -> Dict[str, Any]:
        """Save uploaded brand guideline file"""
        # Create brand-specific upload directory
        brand_upload_dir = self.upload_dir / brand_id
        brand_upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        timestamp = int(time.time())
        file_extension = Path(file.filename).suffix
        safe_filename = f"guideline_{timestamp}{file_extension}"
        file_path = brand_upload_dir / safe_filename
        
        try:
            # Save file
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            return {
                "success": True,
                "filename": safe_filename,
                "original_name": file.filename,
                "size": file_path.stat().st_size,
                "path": str(file_path),
                "upload_time": datetime.now().isoformat()
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

# Initialize service
brand_service = BrandRegistrationService()


def _extract_text_from_guideline(path: Path) -> str:
    """
    Enhanced text extraction from uploaded guideline documents.
    - PDF: Uses PyPDF2 first, falls back to OCR for scanned PDFs
    - DOCX: Uses python-docx
    - Handles scanned documents with OCR (if pytesseract available)
    """
    suffix = path.suffix.lower()
    text_parts = []

    try:
        if suffix == ".pdf" and PdfReader is not None:
            # Try standard PDF text extraction first
            with open(path, "rb") as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
            
            extracted_text = "\n".join(text_parts)
            
            # If we got very little text (likely scanned PDF), try OCR
            if len(extracted_text.strip()) < 100 and OCR_AVAILABLE:
                try:
                    # Convert PDF pages to images
                    images = convert_from_path(str(path), dpi=300)
                    ocr_texts = []
                    for img in images:
                        ocr_text = pytesseract.image_to_string(img)
                        ocr_texts.append(ocr_text)
                    extracted_text = "\n".join(ocr_texts)
                except Exception:
                    # OCR failed, use what we have
                    pass
            
            return extracted_text

        if suffix in {".docx", ".doc"} and docx is not None:
            document = docx.Document(str(path))
            # Extract text from paragraphs
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            # Also try to extract from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
    except Exception as e:
        # Log error but don't break upload flow
        print(f"Text extraction error: {e}")
        return ""

    return ""


def _infer_voice_profile(text: str) -> VoiceProfile:
    """Enhanced voice profile extraction with nuanced detection."""
    lowered = text.lower()
    full_text = text

    # Defaults
    formality = 50
    humor = 40
    warmth = 60

    # Formality detection - more comprehensive
    formality_keywords = {
        "high": ["formal", "professional", "corporate", "executive", "official", "academic", "scholarly"],
        "low": ["casual", "conversational", "relaxed", "informal", "friendly", "laid-back", "chatty"]
    }
    
    formality_score = 0
    for keyword in formality_keywords["high"]:
        if re.search(rf"\b{keyword}\b", lowered):
            formality_score += 1
    for keyword in formality_keywords["low"]:
        if re.search(rf"\b{keyword}\b", lowered):
            formality_score -= 1
    
    if formality_score > 0:
        formality = min(90, 50 + (formality_score * 10))
    elif formality_score < 0:
        formality = max(10, 50 + (formality_score * 10))
    
    # Look for explicit formality percentages
    formality_match = re.search(r"formality[:\s]+(\d{1,2})", lowered)
    if formality_match:
        formality = int(formality_match.group(1))

    # Humor detection - more nuanced
    humor_keywords = {
        "low": ["serious", "no jokes", "no humor", "professional tone", "straightforward", "no-nonsense"],
        "high": ["playful", "fun", "humorous", "quirky", "witty", "lighthearted", "entertaining"]
    }
    
    humor_score = 0
    for keyword in humor_keywords["low"]:
        if re.search(rf"\b{keyword}\b", lowered):
            humor_score -= 1
    for keyword in humor_keywords["high"]:
        if re.search(rf"\b{keyword}\b", lowered):
            humor_score += 1
    
    if humor_score > 0:
        humor = min(90, 40 + (humor_score * 15))
    elif humor_score < 0:
        humor = max(10, 40 + (humor_score * 15))
    
    # Look for explicit humor percentages
    humor_match = re.search(r"humor[:\s]+(\d{1,2})", lowered)
    if humor_match:
        humor = int(humor_match.group(1))

    # Warmth detection - enhanced
    warmth_keywords = {
        "high": ["warm", "friendly", "approachable", "empathetic", "compassionate", "caring", "personal"],
        "low": ["clinical", "impersonal", "corporate", "distant", "detached", "professional"]
    }
    
    warmth_score = 0
    for keyword in warmth_keywords["high"]:
        if re.search(rf"\b{keyword}\b", lowered):
            warmth_score += 1
    for keyword in warmth_keywords["low"]:
        if re.search(rf"\b{keyword}\b", lowered):
            warmth_score -= 1
    
    if warmth_score > 0:
        warmth = min(95, 60 + (warmth_score * 10))
    elif warmth_score < 0:
        warmth = max(20, 60 + (warmth_score * 10))
    
    # Look for explicit warmth percentages
    warmth_match = re.search(r"warmth[:\s]+(\d{1,2})", lowered)
    if warmth_match:
        warmth = int(warmth_match.group(1))

    # Emoji policy - more comprehensive detection
    emoji_policy = "medium"
    emoji_patterns = {
        "none": [r"no emojis", r"avoid emojis", r"without emojis", r"do not use emojis", r"never.*emoji"],
        "light": [r"sparingly.*emoji", r"light use of emojis", r"minimal.*emoji", r"occasional.*emoji"],
        "rich": [r"emoji[- ]rich", r"heavy use of emojis", r"use emojis", r"emoji.*encouraged"]
    }
    
    for policy_level, patterns in emoji_patterns.items():
        for pattern in patterns:
            if re.search(pattern, lowered, re.IGNORECASE):
                emoji_policy = policy_level
                break
        if emoji_policy != "medium":
            break

    return VoiceProfile(
        formality=formality,
        humor=humor,
        warmth=warmth,
        emojiPolicy=emoji_policy,
    )


def _infer_product_share(text: str, default_pct: int = 30) -> int:
    """
    Try to infer 'what share of posts should include the product' from text,
    e.g. '70% of posts can be product-led'.
    """
    match = re.search(r"(\d{1,2})\s*%[^.\n]{0,80}\bproduct", text, flags=re.IGNORECASE)
    if match:
        try:
            value = int(match.group(1))
            if 0 <= value <= 100:
                return value
        except ValueError:
            pass
    return default_pct


def _infer_pillars(text: str) -> list[ContentPillar]:
    """
    Enhanced pillar extraction from various document structures:
    - Bullet points (-, •, *, ·)
    - Numbered lists (1., 2., etc.)
    - Section headings followed by descriptions
    - Content pillars section explicitly labeled
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    pillars: list[ContentPillar] = []
    
    # Look for explicit "Content Pillars" or "Pillars" section
    pillar_section_start = None
    for i, line in enumerate(lines):
        if re.search(r"content\s+pillars?|pillars?", line, re.IGNORECASE):
            pillar_section_start = i
            break
    
    # Extract from pillar section if found
    if pillar_section_start is not None:
        section_lines = lines[pillar_section_start:pillar_section_start + 20]
        for line in section_lines:
            # Skip section header
            if re.search(r"content\s+pillars?|pillars?", line, re.IGNORECASE):
                continue
            
            # Check for bullets or numbers
            if re.match(r"^[-•*·]\s+", line) or re.match(r"^\d+[.)]\s+", line):
                clean = re.sub(r"^[-•*·\d.)]\s+", "", line).strip()
                if len(clean) > 3:
                    name, description = _parse_pillar_line(clean)
                    if name:
                        pillars.append(ContentPillar(name=name, description=description, weight=0))
    
    # If no explicit section found, look for bullet points throughout
    if not pillars:
        bullet_lines = [
            ln for ln in lines 
            if re.match(r"^[-•*·]\s+", ln) or re.match(r"^\d+[.)]\s+", ln)
        ]
        
        for ln in bullet_lines:
            clean = re.sub(r"^[-•*·\d.)]\s+", "", ln).strip()
            if len(clean) < 3:
                continue
            
            name, description = _parse_pillar_line(clean)
            if name and len(name) >= 3:
                # Avoid duplicates
                if not any(p.name.lower() == name.lower() for p in pillars):
                    pillars.append(ContentPillar(name=name, description=description, weight=0))
    
    # Also look for section headings (all caps or title case) as potential pillars
    if len(pillars) < 3:
        for i, line in enumerate(lines):
            # Check if line looks like a heading (short, title case or all caps)
            if (len(line) < 50 and 
                (line.isupper() or line.istitle()) and
                len(line.split()) <= 5):
                # Check if next line(s) contain description
                description = ""
                if i + 1 < len(lines):
                    next_line = lines[i + 1]
                    if len(next_line) > 20 and not next_line[0].isupper():
                        description = next_line[:200]  # Limit description length
                
                name = line.strip()
                if len(name) >= 3 and not any(p.name.lower() == name.lower() for p in pillars):
                    pillars.append(ContentPillar(name=name, description=description, weight=0))
    
    # Calculate weights if we have pillars
    if pillars:
        equal_weight = 100 // len(pillars)
        remainder = 100 % len(pillars)
        for i, pillar in enumerate(pillars):
            pillar.weight = equal_weight + (1 if i < remainder else 0)
    
    return pillars[:10]  # Limit to 10 pillars


def _parse_pillar_line(line: str) -> tuple[str, str]:
    """Parse a single line into pillar name and description."""
    name = line
    description = ""
    
    # Split on various separators
    separators = [" - ", " – ", " — ", ": ", " | ", " — "]
    for sep in separators:
        if sep in line:
            parts = line.split(sep, 1)
            name = parts[0].strip()
            description = parts[1].strip() if len(parts) > 1 else ""
            break
    
    # If no separator, check if line is too long (might be description)
    if not description and len(line) > 60:
        # Try to split at first sentence
        sentence_end = re.search(r"[.!?]\s+", line)
        if sentence_end:
            name = line[:sentence_end.end()].strip()
            description = line[sentence_end.end():].strip()
        else:
            # Split at first comma if very long
            if len(line) > 100 and "," in line:
                parts = line.split(",", 1)
                name = parts[0].strip()
                description = parts[1].strip()
    
    return (name, description)


def _infer_policies(text: str) -> BrandPolicies:
    """
    Enhanced policy extraction:
    - Forbidden phrases from multiple patterns
    - Brand hashtags extraction
    - Max hashtags from document
    """
    lowered = text.lower()
    forbidden: list[str] = []

    # Pattern 1: Quoted phrases after guardrail verbs
    guardrail_patterns = [
        r"(avoid|do not|never say|don't use|prohibited|forbidden)[^\"'\n]*[\"']([^\"']+)[\"']",
        r"[\"']([^\"']+)[\"'][^.\n]*(avoid|do not|never say|prohibited|forbidden)",
    ]
    
    for pattern in guardrail_patterns:
        matches = re.finditer(pattern, lowered, flags=re.IGNORECASE)
        for m in matches:
            phrase = m.group(2) if m.lastindex >= 2 else m.group(1)
            if phrase:
                phrase = phrase.strip()
                if len(phrase) > 2 and phrase.lower() not in [f.lower() for f in forbidden]:
                    forbidden.append(phrase)

    # Pattern 2: Bullet points under "Don't" or "Avoid" sections
    lines = text.splitlines()
    in_forbidden_section = False
    for i, line in enumerate(lines):
        if re.search(r"^(don'?t|avoid|never|prohibited|forbidden)", line, re.IGNORECASE):
            in_forbidden_section = True
            continue
        
        if in_forbidden_section:
            # Check if line is a bullet point
            if re.match(r"^[-•*·]\s+", line.strip()):
                phrase = re.sub(r"^[-•*·]\s+", "", line.strip())
                # Remove quotes if present
                phrase = re.sub(r"^[\"']|[\"']$", "", phrase).strip()
                if len(phrase) > 2 and phrase.lower() not in [f.lower() for f in forbidden]:
                    forbidden.append(phrase)
            elif not line.strip() or re.match(r"^[A-Z]", line.strip()):
                # End of section
                in_forbidden_section = False

    # Pattern 3: Common marketing forbidden phrases if mentioned
    common_forbidden = [
        "buy now", "limited time", "act fast", "click here", 
        "guaranteed", "100% free", "no risk", "urgent"
    ]
    for phrase in common_forbidden:
        if phrase in lowered and phrase not in forbidden:
            # Check if it's in a "don't use" context
            context = lowered[max(0, lowered.find(phrase) - 50):lowered.find(phrase) + len(phrase) + 50]
            if re.search(r"avoid|don'?t|never|prohibited", context, re.IGNORECASE):
                forbidden.append(phrase)

    # Brand hashtags - extract all hashtags
    hashtags = []
    hashtag_matches = re.findall(r"#\w+", text)
    for tag in hashtag_matches:
        if len(tag) > 1 and tag.lower() not in [h.lower() for h in hashtags]:
            hashtags.append(tag)
    
    # Also look for hashtag sections
    hashtag_section_pattern = r"(hashtags?|tags?)[:\s]+([#\w\s,]+)"
    section_matches = re.finditer(hashtag_section_pattern, text, re.IGNORECASE)
    for m in section_matches:
        hashtag_text = m.group(2)
        found_tags = re.findall(r"#\w+", hashtag_text)
        for tag in found_tags:
            if tag.lower() not in [h.lower() for h in hashtags]:
                hashtags.append(tag)

    # Max hashtags extraction
    max_hashtags = 5  # Default
    max_hashtag_patterns = [
        r"max(?:imum)?\s*(?:of\s*)?(\d+)\s*hashtags?",
        r"(\d+)\s*hashtags?\s*(?:max|maximum|limit)",
        r"use\s*(?:up\s*to\s*)?(\d+)\s*hashtags?",
    ]
    
    for pattern in max_hashtag_patterns:
        match = re.search(pattern, lowered, re.IGNORECASE)
        if match:
            try:
                value = int(match.group(1))
                if 1 <= value <= 30:
                    max_hashtags = value
                    break
            except ValueError:
                pass

    return BrandPolicies(
        forbiddenPhrases=forbidden[:50],  # Limit to 50 phrases
        maxHashtags=max_hashtags,
        brandHashtags=sorted(hashtags)[:30],  # Limit to 30 hashtags
    )


def parse_guidelines_to_blueprint(
    text: str, existing_blueprint: Optional[Dict[str, Any]] = None
) -> Optional[BrandBlueprint]:
    """
    Convert free-form guideline text into a draft BrandBlueprint.
    This is intentionally heuristic but gives the Blueprint page a useful
    auto-filled starting point that the user can refine.
    """
    if not text or not text.strip():
        return None

    # Start from existing blueprint if we have one, otherwise defaults
    if isinstance(existing_blueprint, dict):
        base = BrandBlueprint(**existing_blueprint)
    else:
        base = BrandBlueprint()

    voice = _infer_voice_profile(text)
    product_pct = _infer_product_share(text, default_pct=base.productDefaultPct)
    pillars = _infer_pillars(text)
    policies = _infer_policies(text)

    # If we didn't infer any pillars, keep existing ones
    if not pillars:
        pillars = base.pillars

    return BrandBlueprint(
        version=base.version,
        status="generated-from-doc",
        voice=voice,
        pillars=pillars,
        policies=policies,
        productDefaultPct=product_pct,
    )

@router.post("/upload-guideline/{brand_id}")
async def upload_brand_guideline(
    brand_id: str,
    file: UploadFile = File(...)
):
    """Upload brand guideline document"""
    # Validate file type
    allowed_extensions = {'.pdf', '.docx', '.doc'}
    file_extension = Path(file.filename).suffix.lower()
    
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail="Only PDF and DOCX files are allowed"
        )
    
    # Check file size (50MB limit)
    if file.size > 50 * 1024 * 1024:
        raise HTTPException(
            status_code=400,
            detail="File size exceeds 50MB limit"
        )
    
    # Save file
    file_info = brand_service.save_uploaded_file(file, brand_id)

    # Best-effort parse of the uploaded document into a draft blueprint
    extracted_text = _extract_text_from_guideline(Path(file_info["path"]))

    existing_blueprint_dict: Optional[Dict[str, Any]] = None
    try:
        existing_brand = brand_service.get_brand(brand_id)
        existing_blueprint_dict = existing_brand.get("blueprint")  # type: ignore[attr-defined]
    except HTTPException as e:
        if e.status_code != 404:
            raise

    generated_blueprint = parse_guidelines_to_blueprint(
        extracted_text, existing_blueprint=existing_blueprint_dict
    )

    # Prepare updates for brand record
    updates: Dict[str, Any] = {
        "guidelineDoc": {
            "name": file_info["original_name"],
            "size": file_info["size"],
            "status": "uploaded",
            "path": file_info["path"],
            "uploadTime": file_info["upload_time"]
        }
    }

    if generated_blueprint is not None:
        updates["blueprint"] = generated_blueprint.dict()

    # Update or create brand data with file + generated blueprint
    try:
        brand_service.update_brand(brand_id, updates)
    except HTTPException as e:
        if e.status_code == 404:
            # Brand doesn't exist, create it
            now = datetime.now().isoformat()
            brand_data = BrandRegistrationData(
                id=brand_id,
                name=f"Brand {brand_id}",
                guidelineDoc=updates["guidelineDoc"],
                blueprint=generated_blueprint or BrandBlueprint(),  # type: ignore[arg-type]
                createdAt=now,
                updatedAt=now
            )
            brand_service.create_brand(brand_data)
        else:
            raise e

    # Return file info plus any generated blueprint so the frontend can
    # immediately update the Blueprint page without an extra API call.
    response_payload = dict(file_info)
    if generated_blueprint is not None:
        response_payload["blueprint"] = generated_blueprint.dict()

    return response_payload

@router.post("/save-settings/{brand_id}")
async def save_brand_settings(
    brand_id: str,
    settings: BrandSettings
):
    """Save brand default settings"""
    try:
        result = brand_service.update_brand(brand_id, {
            "settings": settings.dict()
        })
        return result
    except HTTPException as e:
        if e.status_code == 404:
            # Brand doesn't exist, create it
            now = datetime.now().isoformat()
            brand_data = BrandRegistrationData(
                id=brand_id,
                name=f"Brand {brand_id}",
                settings=settings,
                createdAt=now,
                updatedAt=now
            )
            return brand_service.create_brand(brand_data)
        else:
            raise e

@router.post("/save-blueprint/{brand_id}")
async def save_brand_blueprint(
    brand_id: str,
    blueprint: BrandBlueprint
):
    """Save brand content blueprint"""
    try:
        result = brand_service.update_brand(brand_id, {
            "blueprint": blueprint.dict()
        })
        return result
    except HTTPException as e:
        if e.status_code == 404:
            # Brand doesn't exist, create it
            now = datetime.now().isoformat()
            brand_data = BrandRegistrationData(
                id=brand_id,
                name=f"Brand {brand_id}",
                blueprint=blueprint,
                createdAt=now,
                updatedAt=now
            )
            return brand_service.create_brand(brand_data)
        else:
            raise e

@router.get("/brand/{brand_id}")
async def get_brand_data(brand_id: str):
    """Get brand registration data"""
    return brand_service.get_brand(brand_id)

@router.get("/brands")
async def list_all_brands():
    """List all brand registrations"""
    return brand_service.list_brands()

@router.post("/create-brand")
async def create_new_brand(brand_data: BrandRegistrationData):
    """Create a new brand registration"""
    return brand_service.create_brand(brand_data)

@router.put("/update-brand/{brand_id}")
async def update_brand_data(brand_id: str, updates: Dict[str, Any]):
    """Update brand registration data"""
    return brand_service.update_brand(brand_id, updates)

@router.get("/export-json/{brand_id}")
async def export_brand_json(brand_id: str):
    """Export complete brand data as JSON file"""
    try:
        brand_data = brand_service.get_brand(brand_id)
        
        # Create export data structure
        export_data = {
            "brand": brand_data,
            "exportedAt": datetime.now().isoformat(),
            "exportVersion": "1.0.0",
            "description": "Complete brand content management data export"
        }
        
        # Create filename
        filename = f"brand-content-data-{brand_id}-{datetime.now().strftime('%Y-%m-%d')}.json"
        
        # Return JSON response with download headers
        from fastapi.responses import JSONResponse
        response = JSONResponse(content=export_data)
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        response.headers["Content-Type"] = "application/json"
        
        return response
        
    except HTTPException:
        raise HTTPException(status_code=404, detail="Brand not found")

@router.post("/regenerate-blueprint/{brand_id}")
async def regenerate_blueprint_from_doc(brand_id: str):
    """Re-generate blueprint from the uploaded guideline document"""
    try:
        brand_data = brand_service.get_brand(brand_id)
    except HTTPException as e:
        if e.status_code == 404:
            raise HTTPException(status_code=404, detail="Brand not found")
        raise e
    
    # Get the uploaded document path
    guideline_doc = brand_data.get("guidelineDoc", {})
    doc_path = guideline_doc.get("path")
    
    if not doc_path or not Path(doc_path).exists():
        raise HTTPException(
            status_code=400,
            detail="No guideline document found. Please upload a document first."
        )
    
    # Extract text from the document
    extracted_text = _extract_text_from_guideline(Path(doc_path))
    
    if not extracted_text or len(extracted_text.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="Could not extract sufficient text from the document. The document may be scanned or corrupted."
        )
    
    # Get existing blueprint to preserve some settings
    existing_blueprint = brand_data.get("blueprint")
    
    # Generate new blueprint from document
    generated_blueprint = parse_guidelines_to_blueprint(
        extracted_text, existing_blueprint=existing_blueprint
    )
    
    if generated_blueprint is None:
        raise HTTPException(
            status_code=500,
            detail="Failed to generate blueprint from document"
        )
    
    # Update blueprint status
    generated_blueprint.status = "regenerated-from-doc"
    
    # Update brand data
    result = brand_service.update_brand(brand_id, {
        "blueprint": generated_blueprint.dict(),
        "updatedAt": datetime.now().isoformat()
    })
    
    return {
        "success": True,
        "message": "Blueprint regenerated successfully from document",
        "blueprint": generated_blueprint.dict()
    }


@router.delete("/brand/{brand_id}")
async def delete_brand(brand_id: str):
    """Delete brand registration"""
    data = brand_service._load_data()
    
    if brand_id not in data["brands"]:
        raise HTTPException(status_code=404, detail="Brand not found")
    
    # Remove brand data
    del data["brands"][brand_id]
    brand_service._save_data(data)
    
    # Clean up uploaded files
    brand_upload_dir = brand_service.upload_dir / brand_id
    if brand_upload_dir.exists():
        shutil.rmtree(brand_upload_dir)
    
    return {"success": True, "message": f"Brand {brand_id} deleted successfully"}